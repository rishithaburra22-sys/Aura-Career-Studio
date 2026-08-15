"""
Aura Career Studio - Utility Functions & Multi-Format Resume Extractors
========================================================================
Supports universal resume extraction from PDF, Word (.docx), Markdown, Text, and JSON.
"""

from __future__ import annotations

import io
import json
import logging
from typing import Dict, Any, Optional

import pypdf

try:
    import docx
except ImportError:
    docx = None

logger = logging.getLogger("aura_career_utils")
logging.basicConfig(level=logging.INFO)


def extract_universal_resume(file_bytes: bytes, filename: str = "resume.pdf") -> str:
    """
    Universally extracts raw text from PDF, DOCX, Markdown, Text, or JSON resume files.
    """
    if not file_bytes:
        return ""

    ext = filename.split(".")[-1].lower() if "." in filename else "pdf"

    # 1. PDF Extraction
    if ext == "pdf":
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes), strict=False)
            extracted_pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    extracted_pages.append(text.strip())
            return "\n\n".join(extracted_pages).strip()
        except Exception as e:
            logger.warning(f"pypdf extraction error for {filename}: {e}")
            try:
                # Fallback to raw decoded text
                return file_bytes.decode("utf-8", errors="ignore").strip()
            except Exception:
                return ""

    # 2. Microsoft Word (.docx) Extraction
    elif ext in ("docx", "doc"):
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            return "\n\n".join(paragraphs).strip()
        except Exception as e:
            logger.warning(f"docx extraction error for {filename}: {e}")
            return file_bytes.decode("utf-8", errors="ignore").strip()

    # 3. JSON Resume Schema
    elif ext in ("json", "jsonl"):
        try:
            data = json.loads(file_bytes.decode("utf-8", errors="ignore"))
            if isinstance(data, dict):
                return json.dumps(data, indent=2)
            return str(data)
        except Exception:
            return file_bytes.decode("utf-8", errors="ignore").strip()

    # 4. Text / Markdown / Code Fallback
    else:
        try:
            return file_bytes.decode("utf-8", errors="ignore").strip()
        except Exception:
            return str(file_bytes)


SAMPLE_RESUMES: Dict[str, Dict[str, Any]] = {
    "fullstack": {
        "title": "Senior Full-Stack Engineer",
        "target_role": "Senior Full-Stack Engineer",
        "github_username": "torvalds",
        "text": """ALEXANDER RIVERS
San Francisco, CA | alex.rivers@example.com | github.com/alexrivers-dev | linkedin.com/in/alex-rivers

PROFESSIONAL SUMMARY
Results-driven Full-Stack Software Engineer with 5+ years of experience designing scalable microservices, reactive user interfaces, and cloud-native systems. Specialized in TypeScript, React, Next.js, Node.js, Python, PostgreSQL, and AWS.

TECHNICAL SKILLS
- Languages: TypeScript, JavaScript (ES6+), Python, SQL, Go (Intermediate)
- Frontend: React 18, Next.js 14, Redux Toolkit, TailwindCSS, WebSockets, HTML5/CSS3
- Backend & Cloud: Node.js, Express, FastAPI, PostgreSQL, Redis, Docker, Kubernetes, AWS (S3, ECS, Lambda)
- Tools: Git, GitHub Actions, Jest, Cypress, Terraform, Datadog

PROFESSIONAL EXPERIENCE
Senior Full-Stack Engineer | CloudScale Systems (2022 - Present)
- Architected a distributed telemetry analytics dashboard using Next.js, TypeScript, and FastAPI, processing 1.2M events/day.
- Reduced API latency by 42% by implementing Redis caching layers and optimizing PostgreSQL indexed queries.
- Spearheaded CI/CD automation with GitHub Actions, cutting release deployment cycle from 45 minutes to 8 minutes.
- Mentored 4 junior engineers on clean architecture patterns, code reviews, and test-driven development.

Full-Stack Developer | Nexus Labs (2020 - 2022)
- Built enterprise customer portal using React and Node.js microservices, supporting 85,000 active monthly users.
- Designed real-time collaborative workspace with WebSockets, decreasing state synchronization latency by 35%.
- Implemented automated end-to-end testing with Cypress, raising overall test coverage from 62% to 91%.

EDUCATION & CERTIFICATIONS
- B.S. in Computer Science | University of California, Berkeley (2016 - 2020)
- AWS Certified Solutions Architect - Associate (2023)
"""
    },
    "aiml": {
        "title": "Staff AI / ML Engineer",
        "target_role": "Staff AI / ML Engineer",
        "github_username": "karpathy",
        "text": """DR. PRIYA SHARMA
Seattle, WA | priya.sharma@example.com | github.com/priyasharma-ai | linkedin.com/in/priya-sharma-ai

PROFESSIONAL SUMMARY
Senior AI/ML Systems Engineer with 6+ years of research and production experience in Large Language Models (LLMs), RAG systems, Vector Search, and distributed model inference. Track record of deploying high-throughput LLM pipelines at scale.

TECHNICAL SKILLS
- AI / ML: PyTorch, Hugging Face Transformers, LangChain, LlamaIndex, vLLM, DeepSpeed, ONNX Runtime
- Vector DBs & Search: FAISS, Qdrant, Pinecone, Milvus, ChromaDB, Cosine Indexing
- Backend & Cloud: Python 3.11, FastAPI, Ray, Triton Inference Server, Docker, Kubernetes, AWS SageMaker
- Data Engineering: PySpark, Pandas, NumPy, PostgreSQL, Redis, Kafka

PROFESSIONAL EXPERIENCE
Lead AI Systems Engineer | HyperScale Intelligence (2022 - Present)
- Designed and deployed multi-tenant Retrieval-Augmented Generation (RAG) agent serving 500,000+ daily queries with sub-80ms retrieval latency.
- Optimized 70B parameter LLM inference using vLLM and Tensor Parallelism, reducing GPU infrastructure spend by 58% ($320k/year).
- Built automated evaluation pipeline assessing hallucination rate, faithfulness, and citation recall across 15,000 benchmark queries.

Machine Learning Engineer | Apex Data Labs (2019 - 2022)
- Developed contextual semantic search engine over 20M enterprise documents using custom fine-tuned BERT embeddings and FAISS.
- Implemented real-time anomaly detection pipeline on Apache Spark handling 50MB/s telemetry data streams.

EDUCATION
- M.S. & B.S. in Artificial Intelligence & Computer Engineering | Stanford University (2014 - 2019)
"""
    },
    "devops": {
        "title": "Cloud DevOps & SRE Lead",
        "target_role": "Cloud DevOps / SRE Lead",
        "github_username": "kelseyhightower",
        "text": """MARCUS VANCE
Austin, TX | marcus.vance@example.com | github.com/marcusvance-ops | linkedin.com/in/marcus-vance-ops

PROFESSIONAL SUMMARY
DevOps & Site Reliability Engineer with 7+ years of experience managing Kubernetes clusters, multi-region cloud infrastructure, and zero-downtime deployment pipelines across AWS and GCP.

TECHNICAL SKILLS
- Cloud Platforms: AWS (EKS, VPC, CloudFront, RDS, IAM), GCP (GKE, Cloud Run)
- Infrastructure as Code: Terraform, Terragrunt, Ansible, CloudFormation
- Containerization: Kubernetes (K8s), Docker, Helm, ArgoCD, Istio Service Mesh
- Observability: Prometheus, Grafana, OpenTelemetry, Datadog, ELK Stack
- CI/CD: GitHub Actions, GitLab CI, Jenkins, Spinnaker

PROFESSIONAL EXPERIENCE
Staff SRE / Platform Lead | Horizon Cloud Inc. (2021 - Present)
- Managed 14 multi-tenant Kubernetes (EKS) clusters hosting 250+ microservices with 99.99% uptime SLA.
- Standardized declarative GitOps workflows using ArgoCD and Helm, reducing rollback incidents by 70%.
- Implemented automated autoscaling policies (HPA/KEDA) that cut monthly cloud compute waste by $45,000.

DevOps Engineer | Strata Security (2018 - 2021)
- Built automated Terraform modules for SOC2-compliant AWS infrastructure provisioning.
- Implemented Prometheus alerting and Grafana dashboards for proactive incident resolution.

EDUCATION & CERTIFICATIONS
- B.S. in Information Systems | UT Austin
- Certified Kubernetes Administrator (CKA), HashiCorp Certified Terraform Associate
"""
    }
}
